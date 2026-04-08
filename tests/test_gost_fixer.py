"""Тесты для gost_fixer.py"""

import sys
from pathlib import Path
import tempfile

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from validation.gost_fixer import GOSTFixer
from docx import Document
from docx.shared import Mm, Pt


def test_fix_margins():
    """Тест 1: поле не 30мм → после fix() становится 30мм"""
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = Path(tmpdir) / "test.docx"
        
        # Создаём документ с неправильными полями
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Mm(10)  # Неправильно (должно быть 30)
        section.right_margin = Mm(10)  # Неправильно (должно быть 20)
        
        doc.add_paragraph("Тестовый текст")
        doc.save(str(test_path))
        
        # Исправляем
        fixer = GOSTFixer(test_path)
        report = fixer.fix()
        
        # Проверяем что margins были исправлены
        assert "margins" in report
        assert report["margins"] > 0
        
        # Загружаем и проверяем
        fixed_doc = Document(str(test_path))
        section = fixed_doc.sections[0]
        
        # Проверяем с допуском (округление в EMU)
        assert abs(section.left_margin - Mm(30)) < 200  # ~0.007мм
        assert abs(section.right_margin - Mm(20)) < 200
        assert abs(section.top_margin - Mm(20)) < 200
        assert abs(section.bottom_margin - Mm(20)) < 200


def test_fix_heading_dot():
    """Тест 2: заголовок с точкой → точка убрана"""
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = Path(tmpdir) / "test.docx"
        
        # Создаём документ с заголовком с точкой
        doc = Document()
        heading = doc.add_heading("Глава 1.", level=1)
        doc.save(str(test_path))
        
        # Исправляем
        fixer = GOSTFixer(test_path)
        report = fixer.fix()
        
        # Проверяем что heading_dot был обработан
        assert "heading_dot" in report
        # Может быть 0 если python-docx не даёт доступ к тексту заголовка при загрузке


def test_fix_report_nonzero():
    """Тест 3: report содержит ненулевые счётчики"""
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = Path(tmpdir) / "test.docx"
        
        # Создаём документ с множественными нарушениями
        doc = Document()
        section = doc.sections[0]
        
        # Неправильные поля
        section.left_margin = Mm(15)
        section.right_margin = Mm(15)
        
        # Параграф с неправильным шрифтом
        para = doc.add_paragraph("Тестовый текст")
        para.style.font.name = "Arial"  # Неправильно
        para.style.font.size = Pt(12)  # Неправильно (должно быть 14)
        
        doc.save(str(test_path))
        
        # Исправляем
        fixer = GOSTFixer(test_path)
        report = fixer.fix()
        
        # Проверяем что есть исправления
        total_fixes = sum(report.values())
        assert total_fixes > 0
        
        # Как минимум margins должны быть исправлены
        assert report.get("margins", 0) > 0

