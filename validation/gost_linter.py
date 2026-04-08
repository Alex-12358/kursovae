"""
GOST Linter - проверка документа без исправления
Возвращает список ошибок, предупреждений и оценку 0.0-1.0
"""
import json
import logging
from pathlib import Path
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

logger = logging.getLogger(__name__)

# ГОСТ параметры (из gost_fixer.py)
GOST_MARGIN_LEFT = 30
GOST_MARGIN_RIGHT = 20  # мм (ГОСТ 2.105)
GOST_MARGIN_TOP = 20
GOST_MARGIN_BOTTOM = 20
GOST_FONT_SIZE = 14
GOST_FONT_NAME = "Times New Roman"
GOST_LINE_SPACING = 1.5
GOST_FIRST_LINE_INDENT = 1.25


class GOSTLinter:
    """
    Проверка документа на соответствие ГОСТ без внесения изменений
    Возвращает JSON с ошибками, предупреждениями и оценкой
    """
    
    def __init__(self, doc_path: Path):
        """
        Args:
            doc_path: Путь к DOCX файлу
        """
        self.doc_path = doc_path
        self.doc = Document(str(doc_path))
        self.errors: List[str] = []
        self.warnings: List[str] = []
        logger.info(f"Инициализирован GOSTLinter для {doc_path}")
    
    def lint(self) -> Dict[str, Any]:
        """
        Выполнить проверку документа
        
        Returns:
            {
                "errors": [...],
                "warnings": [...],
                "score": 0.0-1.0,
                "details": {...}
            }
        """
        logger.info("=== СТАРТ GOST LINTER ===")
        
        self.errors = []
        self.warnings = []
        
        # Выполнить все проверки
        self._check_margins()
        self._check_font()
        self._check_line_spacing()
        self._check_paragraph_indent()
        self._check_alignment()
        self._check_heading_dot()
        
        # Рассчитать оценку
        total_issues = len(self.errors) + len(self.warnings) * 0.5
        max_issues = 50  # условная максимальная планка
        score = max(0.0, 1.0 - (total_issues / max_issues))
        
        result = {
            "errors": self.errors,
            "warnings": self.warnings,
            "score": round(score, 2),
            "details": {
                "errors_count": len(self.errors),
                "warnings_count": len(self.warnings),
                "total_issues": int(total_issues)
            }
        }
        
        logger.info(f"=== ЗАВЕРШЕНО: оценка={score:.2f}, ошибок={len(self.errors)}, предупреждений={len(self.warnings)} ===")
        
        return result
    
    def _check_margins(self):
        """Проверка полей страницы"""
        logger.debug("Проверка полей страницы")
        
        for i, section in enumerate(self.doc.sections):
            if section.left_margin != Mm(GOST_MARGIN_LEFT):
                self.errors.append(
                    f"Раздел {i+1}: левое поле {section.left_margin.mm:.1f}мм, ожидается {GOST_MARGIN_LEFT}мм"
                )
            
            if section.right_margin != Mm(GOST_MARGIN_RIGHT):
                self.errors.append(
                    f"Раздел {i+1}: правое поле {section.right_margin.mm:.1f}мм, ожидается {GOST_MARGIN_RIGHT}мм"
                )
            
            if section.top_margin != Mm(GOST_MARGIN_TOP):
                self.errors.append(
                    f"Раздел {i+1}: верхнее поле {section.top_margin.mm:.1f}мм, ожидается {GOST_MARGIN_TOP}мм"
                )
            
            if section.bottom_margin != Mm(GOST_MARGIN_BOTTOM):
                self.errors.append(
                    f"Раздел {i+1}: нижнее поле {section.bottom_margin.mm:.1f}мм, ожидается {GOST_MARGIN_BOTTOM}мм"
                )
    
    def _check_font(self):
        """Проверка шрифта и размера"""
        logger.debug("Проверка шрифта")
        
        wrong_font_paras = 0
        wrong_size_paras = 0
        
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.name and run.font.name != GOST_FONT_NAME:
                    wrong_font_paras += 1
                    break
        
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.size and run.font.size != Pt(GOST_FONT_SIZE):
                    wrong_size_paras += 1
                    break
        
        if wrong_font_paras > 0:
            self.errors.append(
                f"Неверный шрифт в {wrong_font_paras} параграфах, ожидается {GOST_FONT_NAME}"
            )
        
        if wrong_size_paras > 0:
            self.errors.append(
                f"Неверный размер шрифта в {wrong_size_paras} параграфах, ожидается {GOST_FONT_SIZE}pt"
            )
    
    def _check_line_spacing(self):
        """Проверка межстрочного интервала"""
        logger.debug("Проверка межстрочных интервалов")
        
        wrong_spacing = 0
        
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading'):
                continue
            
            pf = para.paragraph_format
            
            if pf.line_spacing != GOST_LINE_SPACING or pf.line_spacing_rule != WD_LINE_SPACING.MULTIPLE:
                wrong_spacing += 1
        
        if wrong_spacing > 0:
            self.warnings.append(
                f"Неверный межстрочный интервал в {wrong_spacing} параграфах, ожидается {GOST_LINE_SPACING}"
            )
    
    def _check_paragraph_indent(self):
        """Проверка абзацных отступов"""
        logger.debug("Проверка абзацных отступов")
        
        wrong_indent = 0
        
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading') or 'TOC' in para.style.name:
                continue
            
            pf = para.paragraph_format
            
            if pf.first_line_indent != Cm(GOST_FIRST_LINE_INDENT):
                wrong_indent += 1
        
        if wrong_indent > 0:
            self.warnings.append(
                f"Неверный абзацный отступ в {wrong_indent} параграфах, ожидается {GOST_FIRST_LINE_INDENT}см"
            )
    
    def _check_alignment(self):
        """Проверка выравнивания"""
        logger.debug("Проверка выравнивания")
        
        wrong_alignment = 0
        
        for para in self.doc.paragraphs:
            pf = para.paragraph_format
            
            # Заголовки должны быть по центру
            if para.style.name.startswith('Heading'):
                if pf.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    wrong_alignment += 1
            # Основной текст — по ширине
            else:
                if pf.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                    wrong_alignment += 1
        
        if wrong_alignment > 0:
            self.warnings.append(
                f"Неверное выравнивание в {wrong_alignment} параграфах"
            )
    
    def _check_heading_dot(self):
        """Проверка точек после заголовков"""
        logger.debug("Проверка точек после заголовков")
        
        headings_with_dots = 0
        
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading'):
                if para.text.strip().endswith('.'):
                    headings_with_dots += 1
        
        if headings_with_dots > 0:
            self.errors.append(
                f"Недопустимые точки после {headings_with_dots} заголовков (ГОСТ 7.32)"
            )
    
    def to_json(self, result: Dict[str, Any]) -> str:
        """
        Конвертировать результат в JSON
        
        Args:
            result: Результат lint()
            
        Returns:
            JSON строка
        """
        return json.dumps(result, ensure_ascii=False, indent=2)
