"""
GOST Fixer - автоматическое исправление документа по ГОСТ 7.32-2017
Исправляет форматирование без изменения контента
"""
import logging
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

logger = logging.getLogger(__name__)

# ГОСТ параметры
GOST_MARGIN_LEFT = 30  # мм
GOST_MARGIN_RIGHT = 20  # мм (ГОСТ 2.105)
GOST_MARGIN_TOP = 20
GOST_MARGIN_BOTTOM = 20
GOST_FONT_SIZE = 14
GOST_FONT_NAME = "Times New Roman"
GOST_LINE_SPACING = 1.5
GOST_FIRST_LINE_INDENT = 1.25  # см


class GOSTFixer:
    """
    Автоматическое исправление форматирования по ГОСТ
    Фиксит: поля, шрифты, заголовки, интервалы, отступы, выравнивание
    """
    
    FIXES = [
        "margins",  # поля страницы
        "font",  # шрифт и размер
        "line_spacing",  # межстрочный интервал
        "paragraph_indent",  # абзацный отступ
        "alignment",  # выравнивание
        "heading_dot"  # точки после заголовков
    ]
    
    def __init__(self, doc_path: Path):
        """
        Args:
            doc_path: Путь к DOCX файлу
        """
        self.doc_path = doc_path
        self.doc = Document(str(doc_path))
        self.fixes_count: Dict[str, int] = {fix: 0 for fix in self.FIXES}
        logger.info(f"Инициализирован GOSTFixer для {doc_path}")
    
    def fix(self) -> Dict[str, int]:
        """
        Применить все исправления
        
        Returns:
            Словарь с количеством исправлений по каждому типу
        """
        logger.info("=== СТАРТ GOST AUTO-FIXER ===")
        
        self._fix_margins()
        self._fix_font()
        self._fix_line_spacing()
        self._fix_paragraph_indent()
        self._fix_alignment()
        self._fix_heading_dot()
        
        # Сохранить изменения
        self.doc.save(str(self.doc_path))
        
        total_fixes = sum(self.fixes_count.values())
        logger.info(f"=== ЗАВЕРШЕНО: исправлений={total_fixes} ===")
        logger.info(f"Детализация: {self.fixes_count}")
        
        return self.fixes_count
    
    def _fix_margins(self):
        """
        Исправить поля страницы
        ГОСТ 7.32: левое 30мм, правое 15мм, верхнее/нижнее 20мм
        """
        logger.info("Исправление полей страницы")
        
        for section in self.doc.sections:
            changed = False
            
            if section.left_margin != Mm(GOST_MARGIN_LEFT):
                section.left_margin = Mm(GOST_MARGIN_LEFT)
                changed = True
            
            if section.right_margin != Mm(GOST_MARGIN_RIGHT):
                section.right_margin = Mm(GOST_MARGIN_RIGHT)
                changed = True
            
            if section.top_margin != Mm(GOST_MARGIN_TOP):
                section.top_margin = Mm(GOST_MARGIN_TOP)
                changed = True
            
            if section.bottom_margin != Mm(GOST_MARGIN_BOTTOM):
                section.bottom_margin = Mm(GOST_MARGIN_BOTTOM)
                changed = True
            
            if changed:
                self.fixes_count["margins"] += 1
        
        logger.debug(f"Исправлено полей: {self.fixes_count['margins']}")
    
    def _fix_font(self):
        """
        Исправить шрифт и размер
        ГОСТ 7.32: Times New Roman, 14pt
        """
        logger.info("Исправление шрифта")
        
        # Исправить стиль Normal
        style = self.doc.styles['Normal']
        if style.font.name != GOST_FONT_NAME or style.font.size != Pt(GOST_FONT_SIZE):
            style.font.name = GOST_FONT_NAME
            style.font.size = Pt(GOST_FONT_SIZE)
            self.fixes_count["font"] += 1
        
        # Исправить все параграфы
        for para in self.doc.paragraphs:
            for run in para.runs:
                changed = False
                
                if run.font.name != GOST_FONT_NAME:
                    run.font.name = GOST_FONT_NAME
                    changed = True
                
                if run.font.size != Pt(GOST_FONT_SIZE):
                    run.font.size = Pt(GOST_FONT_SIZE)
                    changed = True
                
                if changed:
                    self.fixes_count["font"] += 1
        
        logger.debug(f"Исправлено шрифтов: {self.fixes_count['font']}")
    
    def _fix_line_spacing(self):
        """
        Исправить межстрочный интервал
        ГОСТ 7.32: полуторный (1.5)
        """
        logger.info("Исправление межстрочных интервалов")
        
        for para in self.doc.paragraphs:
            pf = para.paragraph_format
            
            # Проверяем, не является ли параграф заголовком
            if para.style.name.startswith('Heading'):
                continue
            
            changed = False
            
            if pf.line_spacing_rule != WD_LINE_SPACING.MULTIPLE:
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                changed = True
            
            if pf.line_spacing != GOST_LINE_SPACING:
                pf.line_spacing = GOST_LINE_SPACING
                changed = True
            
            if changed:
                self.fixes_count["line_spacing"] += 1
        
        logger.debug(f"Исправлено интервалов: {self.fixes_count['line_spacing']}")
    
    def _fix_paragraph_indent(self):
        """
        Исправить абзацный отступ
        ГОСТ 7.32: первая строка 1.25см
        """
        logger.info("Исправление абзацных отступов")
        
        for para in self.doc.paragraphs:
            # Пропускаем заголовки и оглавление
            if para.style.name.startswith('Heading') or 'TOC' in para.style.name:
                continue
            
            pf = para.paragraph_format
            
            if pf.first_line_indent != Cm(GOST_FIRST_LINE_INDENT):
                pf.first_line_indent = Cm(GOST_FIRST_LINE_INDENT)
                self.fixes_count["paragraph_indent"] += 1
        
        logger.debug(f"Исправлено отступов: {self.fixes_count['paragraph_indent']}")
    
    def _fix_alignment(self):
        """
        Исправить выравнивание
        ГОСТ 7.32: по ширине для основного текста, по центру для заголовков
        """
        logger.info("Исправление выравнивания")
        
        for para in self.doc.paragraphs:
            pf = para.paragraph_format
            
            # Заголовки — по центру
            if para.style.name.startswith('Heading'):
                if pf.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.fixes_count["alignment"] += 1
            # Основной текст — по ширине
            else:
                if pf.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    self.fixes_count["alignment"] += 1
        
        logger.debug(f"Исправлено выравниваний: {self.fixes_count['alignment']}")
    
    def _fix_heading_dot(self):
        """
        Удалить точки после заголовков
        ГОСТ 7.32: точка в конце заголовка не ставится
        """
        logger.info("Удаление точек после заголовков")
        
        for para in self.doc.paragraphs:
            if para.style.name.startswith('Heading'):
                text = para.text.strip()
                
                if text.endswith('.'):
                    # Удалить точку
                    para.text = text[:-1]
                    self.fixes_count["heading_dot"] += 1
        
        logger.debug(f"Удалено точек: {self.fixes_count['heading_dot']}")
