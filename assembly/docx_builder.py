"""
DOCX Builder - сборка документа с обработкой маркеров LLM
Python полностью контролирует форматирование, нумерацию, стили ГОСТ
"""
import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Tuple
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# ГОСТ отступы (из config.py, но дублируем для автономности)
GOST_MARGIN_LEFT = 30  # мм
GOST_MARGIN_RIGHT = 20  # мм (ГОСТ 2.105)
GOST_MARGIN_TOP = 20
GOST_MARGIN_BOTTOM = 20
GOST_FONT_SIZE = 14
GOST_FONT_NAME = "Times New Roman"
GOST_LINE_SPACING = 1.5


class DOCXBuilder:
    """
    Сборщик DOCX документа с обработкой маркеров LLM
    Автоматическая нумерация формул, рисунков, таблиц
    """
    
    def __init__(self, output_path: Path):
        self.output_path = output_path
        self.doc = Document()
        self._apply_gost_styles()
        
        # Счётчики для автонумерации
        self.formula_counter = 0
        self.figure_counter = 0
        self.table_counter = 0
        self.chapter_counter = 0
        
        # Оглавление
        self._toc_entries: List[Tuple[int, str, int]] = []  # (level, title, page)
        
        logger.info(f"Инициализирован DOCXBuilder: {output_path}")
    
    def _apply_gost_styles(self):
        """
        Применить стили ГОСТ к документу
        Поля, шрифты, интервалы согласно ГОСТ 7.32-2017
        """
        sections = self.doc.sections
        for section in sections:
            # Поля страницы
            section.top_margin = Mm(GOST_MARGIN_TOP)
            section.bottom_margin = Mm(GOST_MARGIN_BOTTOM)
            section.left_margin = Mm(GOST_MARGIN_LEFT)
            section.right_margin = Mm(GOST_MARGIN_RIGHT)
        
        # Стиль Normal
        style = self.doc.styles['Normal']
        font = style.font
        font.name = GOST_FONT_NAME
        font.size = Pt(GOST_FONT_SIZE)
        
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = GOST_LINE_SPACING
        paragraph_format.first_line_indent = Cm(1.25)  # абзацный отступ
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        
        logger.debug("Применены стили ГОСТ к документу")
    
    def add_chapter(self, level: int, title: str, content: str):
        """
        Добавить главу с обработкой маркеров LLM
        
        Args:
            level: Уровень заголовка (1, 2, 3)
            title: Название главы
            content: Текст с маркерами [FORMULA], [FIGURE], [TABLE]
        """
        logger.info(f"Добавление главы: {title} (уровень {level})")
        
        # Заголовок
        if level == 1:
            self.chapter_counter += 1
            heading_text = f"{self.chapter_counter} {title}"
        else:
            heading_text = title
        
        heading = self.doc.add_heading(heading_text, level=level)
        heading.style.font.name = GOST_FONT_NAME
        heading.style.font.size = Pt(GOST_FONT_SIZE)
        heading.style.font.bold = True
        
        # Добавить в оглавление
        self._toc_entries.append((level, heading_text, 0))  # page=0 пока заглушка
        
        # КРИТИЧЕСКАЯ ОЧИСТКА: удалить thinking теги перед обработкой маркеров
        content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL | re.IGNORECASE)
        content = content.replace("<think>", "").replace("</think>", "")
        content = content.replace("<THINK>", "").replace("</THINK>", "")
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)  # Очистить лишние переносы
        content = content.strip()
        
        if 'think' in content.lower():
            logger.warning(f"⚠️ DOCX_BUILDER: слово 'think' найдено в главе '{title}'")
        
        # Обработка контента с маркерами
        processed_content = self._process_formulas(content)
        processed_content = self._process_figures(processed_content)
        processed_content = self._process_tables(processed_content)
        
        # Разбить на параграфы
        paragraphs = processed_content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = self.doc.add_paragraph(para_text.strip())
                para.style = self.doc.styles['Normal']
        
        logger.debug(f"Глава '{title}' добавлена")
    
    def _process_formulas(self, text: str) -> str:
        """
        Обработать маркеры [FORMULA]...[/FORMULA]
        Заменить на (N.M) где N - номер главы, M - номер формулы в главе
        
        Args:
            text: Текст с маркерами
            
        Returns:
            Текст с замененными маркерами
        """
        pattern = r'\[FORMULA\](.*?)\[/FORMULA\]'
        
        def replace_formula(match):
            formula_text = match.group(1).strip()
            self.formula_counter += 1
            formula_num = f"({self.chapter_counter}.{self.formula_counter})"
            logger.debug(f"Формула {formula_num}: {formula_text[:50]}...")
            return f"{formula_text} {formula_num}"
        
        result = re.sub(pattern, replace_formula, text, flags=re.DOTALL)
        return result
    
    def _process_figures(self, text: str) -> str:
        """
        Обработать маркеры [FIGURE:описание]
        Заменить на "Рисунок N.M — Описание"
        
        Args:
            text: Текст с маркерами
            
        Returns:
            Текст с замененными маркерами
        """
        pattern = r'\[FIGURE:(.*?)\]'
        
        def replace_figure(match):
            description = match.group(1).strip()
            self.figure_counter += 1
            figure_ref = f"Рисунок {self.chapter_counter}.{self.figure_counter} — {description}"
            logger.debug(f"Рисунок {self.chapter_counter}.{self.figure_counter}: {description}")
            return figure_ref
        
        result = re.sub(pattern, replace_figure, text)
        return result
    
    def _process_tables(self, text: str) -> str:
        """
        Обработать маркеры [TABLE:название]
        Заменить на "Таблица N.M — Название"
        
        Args:
            text: Текст с маркерами
            
        Returns:
            Текст с замененными маркерами
        """
        pattern = r'\[TABLE:(.*?)\]'
        
        def replace_table(match):
            title = match.group(1).strip()
            self.table_counter += 1
            table_ref = f"Таблица {self.chapter_counter}.{self.table_counter} — {title}"
            logger.debug(f"Таблица {self.chapter_counter}.{self.table_counter}: {title}")
            return table_ref
        
        result = re.sub(pattern, replace_table, text)
        return result
    
    def add_toc(self):
        """
        Добавить оглавление в начало документа
        Вызывается из toc_builder.py
        """
        logger.info("Добавление оглавления")
        
        # Вставить разрыв страницы перед оглавлением
        self.doc.add_page_break()
        
        # Заголовок "Содержание"
        toc_heading = self.doc.add_heading("Содержание", level=1)
        toc_heading.style.font.name = GOST_FONT_NAME
        toc_heading.style.font.size = Pt(GOST_FONT_SIZE)
        
        # Добавить записи оглавления
        for level, title, page in self._toc_entries:
            indent = "  " * (level - 1)
            dots = "." * (60 - len(indent) - len(title) - len(str(page)))
            toc_line = f"{indent}{title} {dots} {page}"
            
            para = self.doc.add_paragraph(toc_line)
            para.style = self.doc.styles['Normal']
            para.paragraph_format.first_line_indent = Cm(0)  # без отступа
        
        logger.debug(f"Оглавление добавлено, записей: {len(self._toc_entries)}")
    
    def save(self):
        """
        Сохранить документ
        """
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(self.output_path))
        logger.info(f"Документ сохранён: {self.output_path}")
        logger.info(
            f"Статистика: глав={self.chapter_counter}, "
            f"формул={self.formula_counter}, "
            f"рисунков={self.figure_counter}, "
            f"таблиц={self.table_counter}"
        )
    
    def get_toc_entries(self) -> List[Tuple[int, str, int]]:
        """
        Получить записи оглавления для TOCBuilder
        
        Returns:
            Список (level, title, page)
        """
        return self._toc_entries
    
    def add_drawing_reference(self, name: str, path: str) -> None:
        """
        Добавить ссылку на чертёж в документ.
        
        Args:
            name: Название чертежа (shaft, gear, assembly)
            path: Путь к .dxf файлу
        """
        drawing_names = {
            "shaft": "Чертёж вала",
            "gear": "Чертёж зубчатого колеса", 
            "assembly": "Сборочный чертёж"
        }
        
        title = drawing_names.get(name, f"Чертёж: {name}")
        para = self.doc.add_paragraph()
        para.add_run(f"См. приложение: {title} ({path})")
        para.paragraph_format.first_line_indent = Cm(1.25)
        
        logger.debug(f"Добавлена ссылка на чертёж: {name} -> {path}")
