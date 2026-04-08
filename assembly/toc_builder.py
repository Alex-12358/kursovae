"""
TOC Builder - построение оглавления
Читает записи из DOCXBuilder и форматирует согласно ГОСТ
"""
import logging
from typing import List, Tuple
from docx import Document

logger = logging.getLogger(__name__)


class TOCBuilder:
    """
    Построитель оглавления (содержания)
    Формат: "1 Название главы .............. 5"
    """
    
    def __init__(self, doc: Document, toc_entries: List[Tuple[int, str, int]]):
        """
        Args:
            doc: Объект Document из python-docx
            toc_entries: Список (level, title, page)
        """
        self.doc = doc
        self.toc_entries = toc_entries
        logger.info(f"Инициализирован TOCBuilder с {len(toc_entries)} записями")
    
    def build(self):
        """
        Построить и вставить оглавление в начало документа
        """
        logger.info("Построение оглавления")
        
        # Вставить разрыв страницы в начало
        first_para = self.doc.paragraphs[0]
        first_para._element.addprevious(self.doc.add_page_break()._element)
        
        # Заголовок "Содержание"
        toc_heading = self.doc.add_paragraph("СОДЕРЖАНИЕ")
        toc_heading.style.font.name = "Times New Roman"
        toc_heading.style.font.size = 14
        toc_heading.style.font.bold = True
        toc_heading.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавить записи
        for level, title, page in self.toc_entries:
            self._add_toc_line(level, title, page)
        
        logger.debug("Оглавление построено")
    
    def _add_toc_line(self, level: int, title: str, page: int):
        """
        Добавить строку оглавления
        
        Args:
            level: Уровень заголовка (1, 2, 3)
            title: Текст заголовка
            page: Номер страницы (пока заглушка — 0)
        """
        # Отступ в зависимости от уровня
        indent = "  " * (level - 1)
        
        # Расчёт количества точек для заполнения
        # Длина строки ~60 символов, учитываем отступ, название и номер страницы
        title_with_indent = f"{indent}{title}"
        page_str = str(page) if page > 0 else "0"  # заглушка
        
        # Точки заполнения
        available_space = 60 - len(title_with_indent) - len(page_str) - 2
        dots = "." * max(available_space, 3)
        
        # Формируем строку
        toc_line = f"{title_with_indent} {dots} {page_str}"
        
        # Добавляем параграф
        para = self.doc.add_paragraph(toc_line)
        para.style.font.name = "Times New Roman"
        para.style.font.size = 14
        para.paragraph_format.first_line_indent = 0  # без абзацного отступа
        para.paragraph_format.left_indent = 0
        
        logger.debug(f"Добавлена строка оглавления: {title_with_indent} ... {page_str}")
    
    def update_page_numbers(self, page_mapping: dict):
        """
        Обновить номера страниц в оглавлении
        
        Args:
            page_mapping: Словарь {title: page_number}
        """
        logger.info("Обновление номеров страниц в оглавлении")
        
        for i, (level, title, _) in enumerate(self.toc_entries):
            if title in page_mapping:
                self.toc_entries[i] = (level, title, page_mapping[title])
                logger.debug(f"Обновлена страница для '{title}': {page_mapping[title]}")
